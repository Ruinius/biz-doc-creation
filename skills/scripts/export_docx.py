import sys
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def process_markdown_styles(paragraph, text):
    """Parses text for bold (** or __) and italics (* or _) and adds appropriate runs."""
    # Pattern to find bold and italic segments
    pattern = r'(\*\*(?P<bold1>.*?)\*\*|__(?P<bold2>.*?)__|_(?P<italic1>.*?)_|\*(?P<italic2>.*?)\*)'
    
    last_idx = 0
    for match in re.finditer(pattern, text):
        # Add preceding normal text
        if match.start() > last_idx:
            paragraph.add_run(text[last_idx:match.start()])
        
        # Add styled text
        m_bold = match.group('bold1') or match.group('bold2')
        m_italic = match.group('italic1') or match.group('italic2')
        
        if m_bold:
            run = paragraph.add_run(m_bold)
            run.bold = True
        elif m_italic:
            run = paragraph.add_run(m_italic)
            run.italic = True
        
        last_idx = match.end()
    
    # Add remaining normal text
    if last_idx < len(text):
        paragraph.add_run(text[last_idx:])

def set_cell_padding(cell, top=0, start=0, bottom=0, end=0):
    """Set cell margins (padding) in EMUs. Helper for readable table cells."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:start w:w="{start}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_cell_shading(cell, hex_color):
    """Apply a solid background color to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)

def style_table(table):
    """Apply clean, simple formatting to a completed table: cell padding,
    consistent font sizing, and bold header text. No background shading."""
    CELL_PAD = 80  # ~1.4mm in twips

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_padding(cell, top=CELL_PAD, start=CELL_PAD, bottom=CELL_PAD, end=CELL_PAD)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    if row_idx == 0:
                        run.bold = True

def export_to_docx(markdown_path, docx_path):
    doc = Document()
    
    # Set "Normal" page margins: 1 inch on all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Optional: Customize default styles for a "beautiful" look
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(20)
    h1_style.font.bold = False
    
    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    
    with open(markdown_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    current_table = None

    for line in lines:
        line_stripped = line.rstrip()
        if not line_stripped:
            # End of a table block — apply styling before moving on
            if in_table and current_table:
                style_table(current_table)
            in_table = False
            current_table = None
            continue
            
        leading_spaces = len(line) - len(line.lstrip())
        content = line.strip()

        # Handle Markdown Tables
        if content.startswith('|') and content.endswith('|'):
            cells = [c.strip() for c in content.strip('|').split('|')]
            
            # Skip separator row (e.g., |---|---|)
            if all(all(char in '-: ' for char in c) for c in cells) and len(cells) > 0 and any('-' in c for c in cells):
                continue
                
            if not in_table:
                in_table = True
                current_table = doc.add_table(rows=1, cols=len(cells))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    if i < len(hdr_cells):
                        p = hdr_cells[i].paragraphs[0]
                        process_markdown_styles(p, cell_text)
                        for run in p.runs:
                            run.bold = True
            else:
                row_cells = current_table.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        p = row_cells[i].paragraphs[0]
                        process_markdown_styles(p, cell_text)
            continue
        else:
            in_table = False
            current_table = None

        if content.startswith('# '):
            p = doc.add_heading(level=1)
            process_markdown_styles(p, content[2:])
        elif content.startswith('## '):
            p = doc.add_heading(level=2)
            process_markdown_styles(p, content[3:])
        elif content.startswith('- ') or content.startswith('* '):
            level = leading_spaces // 2
            style_name = 'List Bullet' if level == 0 else f'List Bullet {level + 1}'
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                # Fallback if style doesn't exist
                p = doc.add_paragraph(style='List Bullet')
            process_markdown_styles(p, content[2:])
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            process_markdown_styles(p, content)

    # Style any table that was still open at end of file
    if in_table and current_table:
        style_table(current_table)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python export_docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    export_to_docx(sys.argv[1], sys.argv[2])
